import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Design Palette: Patrimoine Héritage (Automobile & Finance Premium)
COLOR_PRIMARY = RGBColor(10, 14, 26)       # Deep Night Blue (#0A0E1A)
COLOR_SECONDARY = RGBColor(197, 160, 89)   # Sand Gold (#C5A059)
COLOR_TEXT = RGBColor(40, 40, 40)          # Elegant Charcoal (#282828)
COLOR_MUTED = RGBColor(100, 100, 100)      # Gray
HEX_PRIMARY = "0A0E1A"
HEX_SECONDARY = "C5A059"
HEX_LIGHT_BG = "F7F7F9"
HEX_BORDER = "CCCCCC"

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_left_border(cell, color_hex="C5A059", size="24"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    # Left border
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color_hex)
    tcBorders.append(left)
    
    # Clear others
    for border_name in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
        
    tcPr.append(tcBorders)

def add_inline_formatting(paragraph, text):
    # Regex to extract bold, italic, inline code, links
    # Pattern: match bold (**), italic (*), inline code (`)
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.color.rgb = COLOR_PRIMARY
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLOR_PRIMARY
        else:
            # Handle links in text [text](url)
            sub_tokens = re.split(r'(\[.*?\]\(.*?\))', token)
            for sub_token in sub_tokens:
                if sub_token.startswith('[') and '](' in sub_token:
                    match = re.match(r'\[(.*?)\]\((.*?)\)', sub_token)
                    if match:
                        link_text, url = match.groups()
                        run = paragraph.add_run(link_text)
                        run.font.color.rgb = COLOR_SECONDARY
                        run.underline = True
                else:
                    paragraph.add_run(sub_token)

def style_run_default(run, font_name='Calibri', size_pt=11, color_rgb=COLOR_TEXT):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color_rgb

def style_paragraph(p, before_pt=0, after_pt=6, line_spacing=1.15):
    p.paragraph_format.space_before = Pt(before_pt)
    p.paragraph_format.space_after = Pt(after_pt)
    p.paragraph_format.line_spacing = line_spacing

def convert_md_to_docx(md_path, docx_path, title_text):
    doc = Document()
    
    # Configure margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title Block
    title_p = doc.add_paragraph()
    style_paragraph(title_p, before_pt=12, after_pt=6)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title_text)
    style_run_default(run, font_name='Georgia', size_pt=26, color_rgb=COLOR_PRIMARY)
    run.bold = True
    
    subtitle_p = doc.add_paragraph()
    style_paragraph(subtitle_p, before_pt=0, after_pt=24)
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("VeloceWealth · Instrument de Contrôle Financier Automobile")
    style_run_default(sub_run, font_name='Calibri', size_pt=12, color_rgb=COLOR_SECONDARY)
    sub_run.italic = True
    
    # Divider Line
    divider = doc.add_paragraph()
    style_paragraph(divider, before_pt=0, after_pt=18)
    div_run = divider.add_run("—" * 50)
    div_run.font.color.rgb = COLOR_SECONDARY
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    in_mermaid = False
    code_content = []
    
    in_table = False
    table_rows = []
    
    in_list = False
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 1. Fenced Code Blocks & Mermaid
        if stripped.startswith('```'):
            if in_code_block or in_mermaid:
                # Close block
                if in_mermaid:
                    # Represent Mermaid as a beautiful styled box instead of raw code
                    mermaid_text = "\n".join(code_content)
                    
                    p = doc.add_paragraph()
                    style_paragraph(p, before_pt=6, after_pt=12)
                    p.paragraph_format.left_indent = Inches(0.5)
                    
                    table = doc.add_table(rows=1, cols=1)
                    table.autofit = False
                    table.columns[0].width = Inches(5.5)
                    cell = table.cell(0, 0)
                    set_cell_background(cell, HEX_PRIMARY)
                    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
                    set_cell_left_border(cell, color_hex=HEX_SECONDARY, size="36")
                    
                    cp = cell.paragraphs[0]
                    style_paragraph(cp, before_pt=4, after_pt=4)
                    r_head = cp.add_run("📊 SCHÉMA STRATÉGIQUE & PROTOCOLE\n")
                    style_run_default(r_head, font_name='Calibri', size_pt=11, color_rgb=COLOR_SECONDARY)
                    r_head.bold = True
                    
                    # Clean up mermaid syntax for readable representation
                    cleaned_flow = []
                    for m_line in code_content:
                        m_line_strip = m_line.strip()
                        if not m_line_strip or m_line_strip.startswith('graph') or m_line_strip.startswith('style'):
                            continue
                        # Remove mermaid arrow syntax and clean brackets
                        m_line_clean = m_line_strip.replace('-->', ' ➔ ').replace('["', '').replace('"]', '').replace('[', '').replace(']', '').replace('(', '').replace(')', '').replace('<br>', ' ').replace('fill:#0A0E1A,stroke:#C5A059,stroke-width:2px,color:#FDFCF8', '')
                        if m_line_clean.strip():
                            cleaned_flow.append(m_line_clean.strip())
                            
                    r_body = cp.add_run("\n".join(cleaned_flow) if cleaned_flow else mermaid_text)
                    style_run_default(r_body, font_name='Courier New', size_pt=9.5, color_rgb=RGBColor(253, 252, 248))
                    
                else:
                    # Standard Code Block
                    p = doc.add_paragraph()
                    style_paragraph(p, before_pt=6, after_pt=12)
                    p.paragraph_format.left_indent = Inches(0.5)
                    
                    table = doc.add_table(rows=1, cols=1)
                    table.autofit = False
                    table.columns[0].width = Inches(5.5)
                    cell = table.cell(0, 0)
                    set_cell_background(cell, HEX_LIGHT_BG)
                    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                    set_cell_left_border(cell, color_hex="999999", size="24")
                    
                    cp = cell.paragraphs[0]
                    style_paragraph(cp, before_pt=2, after_pt=2)
                    r = cp.add_run("\n".join(code_content))
                    style_run_default(r, font_name='Courier New', size_pt=9.5, color_rgb=COLOR_PRIMARY)
                
                in_code_block = False
                in_mermaid = False
                code_content = []
            else:
                # Open block
                if 'mermaid' in stripped:
                    in_mermaid = True
                else:
                    in_code_block = True
            i += 1
            continue
            
        if in_code_block or in_mermaid:
            code_content.append(line.rstrip('\n'))
            i += 1
            continue

        # 2. Table Parsing
        if stripped.startswith('|'):
            in_table = True
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            # Table ended or empty line
            in_table = False
            # Render the table
            if len(table_rows) >= 2:
                # Filter out separator row (contains only dashes, pipes, colons)
                rows_data = []
                for tr in table_rows:
                    if re.match(r'^\|[\s\-\|:]+\|$', tr):
                        continue
                    # Split cells and strip
                    cells = [c.strip() for c in tr.split('|')[1:-1]]
                    rows_data.append(cells)
                
                if rows_data:
                    num_cols = len(rows_data[0])
                    num_rows = len(rows_data)
                    
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.autofit = True
                    
                    for r_idx, row_cells in enumerate(rows_data):
                        for c_idx, cell_text in enumerate(row_cells):
                            cell = table.cell(r_idx, c_idx)
                            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                            
                            # Styling Table
                            if r_idx == 0:
                                # Header Style
                                set_cell_background(cell, HEX_PRIMARY)
                                p = cell.paragraphs[0]
                                style_paragraph(p, before_pt=4, after_pt=4)
                                run = p.add_run(cell_text)
                                style_run_default(run, font_name='Calibri', size_pt=10.5, color_rgb=RGBColor(253, 252, 248))
                                run.bold = True
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                # Alternating row colors
                                if r_idx % 2 == 1:
                                    set_cell_background(cell, HEX_LIGHT_BG)
                                p = cell.paragraphs[0]
                                style_paragraph(p, before_pt=4, after_pt=4)
                                # Check alignment
                                add_inline_formatting(p, cell_text)
                                for r in p.runs:
                                    style_run_default(r, font_name='Calibri', size_pt=10, color_rgb=COLOR_TEXT)
                                
                                # Center check marks or numeric values
                                if cell_text in ['Oui', 'Non', 'Illimité', '1', '3 offertes', 'Jusqu\'à 5'] or '€' in cell_text or '%' in cell_text:
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                    # Add space after table
                    spacer = doc.add_paragraph()
                    style_paragraph(spacer, before_pt=0, after_pt=12)
            table_rows = []
            # Do not skip current line as it's not a table line

        # 3. Headings
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            style_paragraph(p, before_pt=18, after_pt=8)
            run = p.add_run(stripped[2:])
            style_run_default(run, font_name='Georgia', size_pt=20, color_rgb=COLOR_PRIMARY)
            run.bold = True
            
            # Gold accent line below Heading 1
            border_p = doc.add_paragraph()
            style_paragraph(border_p, before_pt=0, after_pt=12)
            b_run = border_p.add_run("―" * 30)
            b_run.font.color.rgb = COLOR_SECONDARY
            b_run.font.size = Pt(8)
            
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            style_paragraph(p, before_pt=16, after_pt=6)
            run = p.add_run(stripped[3:])
            style_run_default(run, font_name='Georgia', size_pt=15, color_rgb=COLOR_PRIMARY)
            run.bold = True
            
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            style_paragraph(p, before_pt=12, after_pt=4)
            run = p.add_run(stripped[4:])
            style_run_default(run, font_name='Calibri', size_pt=12.5, color_rgb=COLOR_SECONDARY)
            run.bold = True
            
        elif stripped.startswith('#### '):
            p = doc.add_paragraph()
            style_paragraph(p, before_pt=8, after_pt=2)
            run = p.add_run(stripped[5:])
            style_run_default(run, font_name='Calibri', size_pt=11.5, color_rgb=COLOR_PRIMARY)
            run.bold = True
            run.italic = True

        # 4. Blockquotes & Alerts (> )
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            style_paragraph(p, before_pt=6, after_pt=8)
            p.paragraph_format.left_indent = Inches(0.4)
            
            # Use table with single cell to style like alert box
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.columns[0].width = Inches(5.6)
            cell = table.cell(0, 0)
            set_cell_background(cell, HEX_LIGHT_BG)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            
            # Determine color and icon depending on alert type
            border_color = HEX_SECONDARY
            icon = "ℹ️ "
            content = stripped[2:]
            if "[!IMPORTANT]" in content:
                border_color = "E03E2F"
                icon = "🔥 OBLIGATOIRE : "
                content = content.replace("[!IMPORTANT]", "")
            elif "[!WARNING]" in content:
                border_color = "FF9F00"
                icon = "⚠️ ATTENTION : "
                content = content.replace("[!WARNING]", "")
            elif "[!NOTE]" in content:
                border_color = HEX_PRIMARY
                icon = "📝 NOTE : "
                content = content.replace("[!NOTE]", "")
            elif "[!TIP]" in content:
                border_color = "2ECC71"
                icon = "💡 ASTUCE : "
                content = content.replace("[!TIP]", "")
                
            set_cell_left_border(cell, color_hex=border_color, size="36")
            cp = cell.paragraphs[0]
            style_paragraph(cp, before_pt=2, after_pt=2)
            
            # Bold prefix
            run_pref = cp.add_run(icon)
            style_run_default(run_pref, font_name='Calibri', size_pt=10.5, color_rgb=COLOR_PRIMARY)
            run_pref.bold = True
            
            add_inline_formatting(cp, content.strip())
            for r in cp.runs[1:]:
                style_run_default(r, font_name='Calibri', size_pt=10.5, color_rgb=COLOR_TEXT)
                r.italic = True

        # 5. Lists (Bullet & Numbered)
        elif stripped.startswith('* ') or stripped.startswith('- ') or re.match(r'^\d+\.\s', stripped):
            # Check indentation level to determine indent size
            indent_spaces = len(line) - len(line.lstrip())
            indent_inch = 0.25 + (indent_spaces // 2) * 0.25
            
            p = doc.add_paragraph()
            style_paragraph(p, before_pt=2, after_pt=3)
            p.paragraph_format.left_indent = Inches(indent_inch)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            
            if re.match(r'^\d+\.\s', stripped):
                # Numbered list
                match = re.match(r'^(\d+\.)\s(.*)', stripped)
                num_prefix, list_content = match.groups()
                
                run_num = p.add_run(f"{num_prefix}  ")
                style_run_default(run_num, font_name='Calibri', size_pt=11, color_rgb=COLOR_SECONDARY)
                run_num.bold = True
                
                add_inline_formatting(p, list_content)
            else:
                # Bullet list
                bullet_char = "▪  " if indent_spaces > 0 else "•  "
                run_bullet = p.add_run(bullet_char)
                style_run_default(run_bullet, font_name='Calibri', size_pt=11, color_rgb=COLOR_SECONDARY)
                run_bullet.bold = True
                
                # Content
                add_inline_formatting(p, stripped[2:])
                
            for r in p.runs[1:]:
                style_run_default(r, font_name='Calibri', size_pt=11, color_rgb=COLOR_TEXT)

        # 6. Normal Paragraph
        elif stripped:
            p = doc.add_paragraph()
            style_paragraph(p, before_pt=4, after_pt=8)
            add_inline_formatting(p, stripped)
            for r in p.runs:
                style_run_default(r, font_name='Calibri', size_pt=11, color_rgb=COLOR_TEXT)
                
        i += 1
        
    doc.save(docx_path)
    print(f"✓ Document Word généré avec succès : {docx_path}")

if __name__ == "__main__":
    brain_dir = "/Users/jinola/.gemini/antigravity-ide/brain/6582f0ae-98a0-43d1-bee2-902d56c3d8b6"
    
    # 1. Convert Strategy Document
    strategy_md = os.path.join(brain_dir, "launch_and_marketing_strategy.md")
    strategy_docx = os.path.join(brain_dir, "VeloceWealth_Launch_And_Marketing_Strategy.docx")
    if os.path.exists(strategy_md):
        convert_md_to_docx(strategy_md, strategy_docx, "Stratégie de Lancement & Marketing")
        
    # 2. Convert Product Specifications
    prd_md = os.path.join(brain_dir, "product_specifications.md")
    prd_docx = os.path.join(brain_dir, "VeloceWealth_Product_Specifications.docx")
    if os.path.exists(prd_md):
        convert_md_to_docx(prd_md, prd_docx, "Spécifications Fonctionnelles & Produit")
